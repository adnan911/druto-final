from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/home/ubuntu/druto-platform')
DOCS = ROOT / 'docs'
OUT = DOCS / 'ARC_TEAM_PRESENTATION_COMPLETE_SCRIPT.docx'
SOURCES = [
    ('Arc Team Demo Pitch Script and Technical Q&A', DOCS / 'ARC_TEAM_DEMO_PITCH_SCRIPT.md'),
    ('Complete Slide-by-Slide Presentation Narration', DOCS / 'ARC_TEAM_PRESENTATION_FULL_SCRIPT.md'),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def add_rich_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        hdr[i].text = value.strip()
        set_cell_shading(hdr[i], 'DDEFE7')
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value.strip()
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith('|'):
        line = lines[index].strip()
        values = [v.strip() for v in line.strip('|').split('|')]
        if all(re.fullmatch(r':?-{3,}:?', v) for v in values):
            index += 1
            continue
        rows.append(values)
        index += 1
    return rows, index


def add_markdown(doc, title, path):
    doc.add_heading(title, level=1)
    lines = path.read_text(encoding='utf-8').splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph(style='Intense Quote')
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            i += 1
            continue
        if line.startswith('|'):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            add_rich_text(p, line[2:])
        elif re.match(r'^\s*[-*] ', raw):
            p = doc.add_paragraph(style='List Bullet')
            add_rich_text(p, re.sub(r'^\s*[-*] ', '', raw))
        elif re.match(r'^\s*\d+\. ', raw):
            p = doc.add_paragraph(style='List Number')
            add_rich_text(p, re.sub(r'^\s*\d+\. ', '', raw))
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line)
        i += 1


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)
for name, size, color in [('Title', 24, '17302E'), ('Heading 1', 18, '17302E'), ('Heading 2', 14, '1E9B83'), ('Heading 3', 11.5, '2758D9')]:
    style = styles[name]
    style.font.name = 'Aptos Display' if name != 'Heading 3' else 'Aptos'
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DRUTO × ARC')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(30, 155, 131)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Arc Team Demonstration Presentation Script')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(23, 32, 30)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Real Arc Testnet USDC flow · Buyer checkout · Onchain proof · Seller dashboard').italic = True
doc.add_page_break()

for title, path in SOURCES:
    add_markdown(doc, title, path)
    if path != SOURCES[-1][1]:
        doc.add_page_break()

props = doc.core_properties
props.title = 'Druto × Arc — Complete Arc Team Demonstration Presentation Script'
props.author = 'Manus AI'
props.subject = 'Arc Testnet USDC marketplace payment demonstration'
props.keywords = 'Druto, Arc, USDC, presentation, seller dashboard, technical Q&A'
doc.save(OUT)
print(OUT)
